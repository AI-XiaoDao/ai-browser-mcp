# -*- coding: utf-8 -*-
"""Generate AI Browser MCP promo Word document with advanced styling."""
from pathlib import Path
import urllib.request

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "AI-Browser-MCP-Promo-v2.6.docx"
IMG_URL = "https://github.com/AI-XiaoDao/ai-browser-mcp/raw/main/.github/demo-douyin-post-scan.png"
IMG_LOCAL = ROOT / "release" / "_promo_demo.png"

# Brand colors
C_PRIMARY = RGBColor(0x00, 0x66, 0xCC)   # MCP blue
C_ACCENT = RGBColor(0x23, 0x86, 0x36)    # green
C_PURPLE = RGBColor(0x6B, 0x46, 0xC1)
C_TEXT = RGBColor(0x24, 0x29, 0x2F)
C_MUTED = RGBColor(0x58, 0x60, 0x69)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_HEADER_BG = RGBColor(0x00, 0x66, 0xCC)
C_ROW_ALT = RGBColor(0xF0, 0xF6, 0xFC)


def set_cell_shading(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_run_font(run, name="Microsoft YaHei", size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_para(doc, text, style=None, align=None, space_after=6, space_before=0):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=11, color=C_TEXT)
    return p


def add_rich_para(doc, parts, align=None, space_after=8):
    """parts: list of (text, bold, color, size)"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    for text, bold, color, size in parts:
        run = p.add_run(text)
        set_run_font(run, size=size or 11, bold=bold, color=color or C_TEXT)
    return p


def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(8)
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: C_PRIMARY, 2: C_PURPLE, 3: C_TEXT}
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 12), bold=True, color=colors.get(level, C_TEXT))
    # left accent bar via border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), "0066CC")
    left.set(qn("w:space"), "4")
    pBdr.append(left)
    pPr.append(pBdr)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=C_TEXT)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=11, color=C_TEXT)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "E8F4FD")
    p = cell.paragraphs[0]
    r1 = p.add_run(title + "\n")
    set_run_font(r1, size=11, bold=True, color=C_PRIMARY)
    r2 = p.add_run(body)
    set_run_font(r2, size=10.5, color=C_MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "0066CC")
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True, color=C_WHITE)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        if ri % 2 == 1:
            for c in cells:
                set_cell_shading(c, "F0F6FC")
        for ci, val in enumerate(row):
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10, color=C_TEXT)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return table


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "1E1E1E")
    p = cell.paragraphs[0]
    run = p.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xD4, 0xD4, 0xD4)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def setup_styles(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    # header
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("AI Browser MCP Server v2.6 · MIT")
    set_run_font(r, size=8, color=C_MUTED, italic=True)
    # footer
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("https://github.com/AI-XiaoDao/ai-browser-mcp")
    set_run_font(r, size=8, color=C_PRIMARY)


def download_image():
    try:
        urllib.request.urlretrieve(IMG_URL, IMG_LOCAL)
        return IMG_LOCAL.exists()
    except Exception:
        return False


def build():
    doc = Document()
    setup_styles(doc)

    # ── Cover ──
    for _ in range(3):
        doc.add_paragraph()
    add_rich_para(
        doc,
        [("AI Browser MCP Server", True, C_PRIMARY, 28), ("\n", False, C_TEXT, 28)],
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )
    add_rich_para(
        doc,
        [
            ("任意 AI 代理 · 一句话操控 Windows 真实浏览器", False, C_MUTED, 14),
            ("\n", False, C_MUTED, 14),
        ],
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=20,
    )
    badges = "MCP 标准协议  |  217 工具全开放  |  FBrowser CEF  |  MIT 开源  |  Windows x64"
    add_para(doc, badges, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_rich_para(
        doc,
        [
            ("GitHub：", True, C_TEXT, 11),
            ("https://github.com/AI-XiaoDao/ai-browser-mcp", False, C_PRIMARY, 11),
            ("\n", False, C_TEXT, 11),
            ("下载：", True, C_TEXT, 11),
            ("https://github.com/AI-XiaoDao/ai-browser-mcp/releases/tag/v2.6.0", False, C_PRIMARY, 11),
        ],
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=30,
    )
    add_para(doc, "作者 QQ：212577526", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    doc.add_page_break()

    # ── 1 概述 ──
    add_heading_styled(doc, "一、产品概述", 1)
    add_para(
        doc,
        "AI浏览器 MCP Server 是基于火山视窗 + FBrowser CEF 的 Windows 本地浏览器自动化 MCP 服务。"
        "启动 AI浏览器.exe 后，本机 9222 端口对外提供标准 Model Context Protocol (MCP) 接口。"
        "不限于 Cursor — Claude Desktop、Cline、OpenCode、自研 Agent 或 HTTP/WebSocket 脚本均可调用 217 个 browser_* 工具。"
        "说一句话，Agent 自动串联执行，无需手写 Playwright。",
    )
    add_callout(
        doc,
        "🎁 GitHub Release 下载即用",
        "AI-Browser-MCP-x64-v2.6.0.zip（约 157MB）已包含全部 217 个工具（截图、CDP 调试器、指纹、网络拦截、工作流等），解压运行即可。",
    )
    add_para(doc, "架构：AI 代理 → mcp_bridge.js（可选）→ MCP Server → FBrowser CEF → 真实网页", space_after=12)

    add_heading_styled(doc, "接入方式", 2)
    add_table(
        doc,
        ["方式", "适用场景"],
        [
            ["stdio MCP", "Cursor、Claude Desktop、Cline — mcp_bridge.js 桥接"],
            ["HTTP POST", "任意语言 / 自研 Agent — POST http://127.0.0.1:9222/mcp"],
            ["WebSocket", "长连接 JSON-RPC — ws://127.0.0.1:9222"],
        ],
        col_widths=[3.5, 12],
    )

    # ── 2 亮点 ──
    add_heading_styled(doc, "二、核心亮点", 1)
    highlights = [
        ("自然语言驱动", "在任意 MCP AI 代理里说需求，Agent 自动选工具链，逐步 sync-wait 等结果。"),
        ("217 开箱工具", "导航、填表、DOM、JS、网络、截图、CDP 断点、指纹、Hook、工作流等 24 大类。"),
        ("真实 CEF 窗口", "FBrowser 内核，页面行为贴近用户日常浏览，非纯无头模拟。"),
        ("sync-wait + batch", "同次调用内等待结果；一次请求串联多工具，省 Token、少轮次。"),
        ("本机优先", "默认 127.0.0.1:9222，数据不出机器；关主窗口后托盘常驻。"),
        ("MIT 开源", "11 个 .wsv 模块（~2 万行）+ generated-cpp + 文档 + 场景脚本，欢迎 PR。"),
    ]
    for title, desc in highlights:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run(f"▎{title}  ")
        set_run_font(r1, size=11, bold=True, color=C_PRIMARY)
        r2 = p.add_run(desc)
        set_run_font(r2, size=10.5, color=C_MUTED)

    # ── 3 场景 ──
    add_heading_styled(doc, "三、一句话场景", 1)
    add_para(doc, "不用写脚本、不用记 217 个工具名 — 复制话术到 AI 对话即可。", space_after=10)
    add_table(
        doc,
        ["#", "场景", "示例话术", "主要工具"],
        [
            ["①", "数据采集", "滚动列表，采标题价格 JSON", "dom_query / collect"],
            ["②", "逆向分析", "扫描 POST，标疑似加密字段", "inject Hook / network"],
            ["③", "定位算法", "断点跟到 sign 函数，给源码", "debugger_*"],
            ["④", "自动填表", "登录后台，导出订单表", "fill_* / workflow"],
            ["⑤", "工作流复用", "存成 JSON，下次一键跑", "workflow_*"],
        ],
        col_widths=[1.2, 2.5, 5.5, 4.5],
    )

    add_heading_styled(doc, "话术示例（复制到 AI 代理）", 2)
    prompts = [
        "帮我把商品列表滚动 5 屏，把标题、价格、链接采成 JSON。",
        "打开 https://www.douyin.com ，扫描 XHR/fetch 的 POST，标出疑似加密字段并说明依据。",
        "提交表单时下断点，跟到计算 sign 的 JS 函数，把函数名和源码片段给我。",
        "用账号登录后台，进入订单页，导出前 20 行表格；把流程存成 workflow JSON。",
    ]
    for pr in prompts:
        add_bullet(doc, f"「{pr}」")

    add_heading_styled(doc, "以前 vs 现在", 2)
    add_table(
        doc,
        ["任务", "传统 Playwright", "AI浏览器 MCP"],
        [
            ["采商品列表", "写选择器 + 滚动循环 + 解析", "一句话 → Agent 调 MCP"],
            ["找 POST 加密字段", "手动 DevTools Hook", "一句话 + inject Hook"],
            ["找 sign 函数", "手动断点跟栈", "一句话 + debugger_* 编排"],
            ["固定流程", "维护 cron + 脚本", "workflow JSON + workflow_run"],
        ],
        col_widths=[3.5, 5.5, 5.5],
    )

    # Demo image
    add_heading_styled(doc, "实测演示", 2)
    add_para(doc, "Cursor + Agent 一句话：自动注入 Hook、滚动触发请求、分析 33 条 POST。", space_after=8)
    if download_image():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(IMG_LOCAL), width=Inches(5.8))
        cap = doc.add_paragraph("抖音 POST 逆向扫描 — Hook XHR/fetch · 标出疑似加密字段")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cap.runs[0], size=9, italic=True, color=C_MUTED)

    doc.add_page_break()

    # ── 4 扩展 ──
    add_heading_styled(doc, "四、强大扩展能力", 1)
    add_table(
        doc,
        ["扩展", "说明"],
        [
            ["217 MCP 工具", "导航/填表/DOM/JS/网络/截图/CDP/调试器/指纹 … 24 大类"],
            ["工作流 JSON", "workflows/*.json 多步骤编排，workflow_run 一键复跑"],
            ["事件 Hook", "生命周期/网络/资源 persist 注入，逆向与改包"],
            ["场景脚本", "scenarios/ 逆向扫描、Hook 测试、断点恢复范例"],
            ["技能书 + 测试", "skills/ Agent 知识库；run_all_tests.js 全量回归"],
            ["多通道接入", "stdio / HTTP POST / WebSocket 任意 AI 程序"],
        ],
        col_widths=[4, 11.5],
    )

    add_heading_styled(doc, "工具速查", 2)
    add_table(
        doc,
        ["类别", "示例工具", "说明"],
        [
            ["导航", "browser_navigate", "打开 URL、后退、刷新"],
            ["填表", "browser_fill_click", "填值、点击、选择"],
            ["DOM", "browser_dom_query", "查询元素、属性、文本"],
            ["网络", "browser_network", "请求列表、collect 聚合"],
            ["Hook", "browser_inject", "persist 注入抓 POST body"],
            ["调试", "browser_debugger_*", "断点、单步、栈、源码"],
            ["工作流", "workflow_run", "JSON 多步骤一键执行"],
        ],
        col_widths=[2.5, 4.5, 8.5],
    )

    # ── 5 上手 ──
    add_heading_styled(doc, "五、3 步上手", 1)
    add_numbered(doc, "下载 AI-Browser-MCP-x64-v2.6.0.zip（~157MB）→ 解压 → 双击 AI浏览器.exe")
    add_numbered(doc, "浏览器打开 http://127.0.0.1:9222/health ，确认 \"status\":\"ok\"")
    add_numbered(doc, "配置 mcp_bridge.js 接入 Cursor / Claude → 对 Agent 说任一句场景话术")

    add_heading_styled(doc, "Cursor 配置示例", 2)
    add_code_block(
        doc,
        '{\n  "mcpServers": {\n    "ai-browser": {\n      "command": "node",\n      "args": ["路径/mcp_bridge.js"],\n      "env": {\n        "AI_BROWSER_MCP_HTTP_POST": "http://127.0.0.1:9222/mcp"\n      }\n    }\n  }\n}',
    )
    add_para(doc, "Release 解压目录请将 args 改为本机 mcp_bridge.js 路径。自检：node mcp_bridge.js --check", space_after=12)

    add_heading_styled(doc, "环境要求", 2)
    add_bullet(doc, "Windows 10 / 11（64 位）")
    add_bullet(doc, "Node.js 18+（仅 Cursor 桥接需要）")
    add_bullet(doc, "Release 已含 CEF 运行时，无需单独安装浏览器")

    # ── 6 开源 ──
    add_heading_styled(doc, "六、开源与下载", 1)
    add_table(
        doc,
        ["内容", "说明"],
        [
            ["GitHub 仓库", "https://github.com/AI-XiaoDao/ai-browser-mcp"],
            ["Release 运行包", "AI-Browser-MCP-x64-v2.6.0.zip — 217 工具全开放"],
            ["权威源码", "CEFbro/AI浏览器/src/*.wsv（11 模块，MIT）"],
            ["C++ 对照", "generated-cpp/release-x64/"],
            ["文档", "docs/ · skills/ · FORUM_POSTS.md · promo.html"],
        ],
        col_widths=[3.5, 12],
    )

    add_heading_styled(doc, "仓库结构（摘要）", 2)
    add_code_block(
        doc,
        "ai-browser-mcp/\n├── CEFbro/AI浏览器/src/     # .wsv MCP 核心\n├── generated-cpp/             # C++ 对照\n├── docs/ · skills/            # 文档与技能书\n├── mcp_bridge.js              # Cursor 桥接\n├── workflows/ · scenarios/   # 工作流与场景\n└── release/pack-release.ps1  # 发版脚本",
    )

    # ── 7 FAQ ──
    add_heading_styled(doc, "七、常见问题", 1)
    faqs = [
        ("和 Playwright 有什么区别？", "217 工具已封装成 MCP，AI 代理直接调用，不用写 Node 浏览器脚本。真实 FBrowser CEF 窗口，本机 9222。"),
        ("只能 Cursor 吗？", "否。任意 MCP 客户端或 HTTP POST 127.0.0.1:9222/mcp 均可。"),
        ("能抓 POST body 吗？", "默认网络工具不记 POST 正文；须 browser_inject persist Hook。见 douyin_xhr_encrypt_scan.js。"),
        ("Mac / Linux 支持吗？", "当前仅 Windows x64。CEF 运行时随 Release 附带。"),
        ("逆向和定位有什么区别？", "逆向找哪个请求/字段被加密；定位找加密逻辑在哪个 JS 函数（CDP 断点 + 栈 + 源码）。"),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rq = p.add_run(f"Q：{q}\n")
        set_run_font(rq, size=10.5, bold=True, color=C_PRIMARY)
        ra = p.add_run(f"A：{a}")
        set_run_font(ra, size=10.5, color=C_MUTED)
        p.paragraph_format.space_after = Pt(10)

    # ── Footer block ──
    doc.add_paragraph()
    add_callout(
        doc,
        "参与与反馈",
        "欢迎 Star ⭐ · GitHub Issue · Pull Request\n"
        "宣传材料：FORUM_POSTS.md · promo.html · OPEN_SOURCE.md\n"
        "许可证：MIT License",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("作者 QQ：212577526")
    set_run_font(r, size=12, bold=True, color=C_PRIMARY)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
